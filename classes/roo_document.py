import docx
import json
import sys
import os
import csv
import shutil
from dotenv import load_dotenv

from classes.database import Database
from classes.classification import Classification
from classes.table_cell import TableCell, Rule


class RooDocument(object):
    def __init__(self):
        self.get_args()
        self.get_config()
        self.load_exemplar_codes()
        self.get_schemes()

    def get_args(self):
        if len(sys.argv) < 3:
            print("Please specify a source file and the name of the agreement")
            sys.exit()
        else:
            path = os.getcwd()
            self.path_resources = os.path.join(path, "resources")
            self.path_source = os.path.join(self.path_resources, "source")
            self.path_dest = os.path.join(self.path_resources, "dest")

            filename = sys.argv[1]
            self.agreement = sys.argv[2]
            self.source = os.path.join(self.path_source, filename + ".docx")

    def get_config(self):
        # Get variables from the env file
        load_dotenv('.env')
        self.max_row_count = int(os.getenv('MAX_ROW_COUNT'))
        self.scheme_file = os.getenv('SCHEME_FILE')
        self.fetch_descriptions = os.getenv('FETCH_DESCRIPTIONS')
        self.exemplar_codes_path = os.getenv('EXEMPLAR_CODES_PATH')
        self.roo_folder = os.getenv('ROO_FOLDER')
        self.dest_filename = os.path.join(
            self.roo_folder, self.agreement.lower() + "_roo.json")

    def load_exemplar_codes(self):
        # Loads up all of the exemplar codes from the file generated in the 'EU' projectq
        self.exemplar_codes = []
        with open(self.exemplar_codes_path) as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=',')
            line_count = 0
            for row in csv_reader:
                ex = Classification(row[0], row[1], "80", -1, 1)
                self.exemplar_codes.append(ex)

    def get_schemes(self):
        self.rule_offset = 0
        self.country_prefix = ""

        # Get the offset from the schemes file
        f = open(self.scheme_file,)
        data = json.load(f)
        for item in data["schemes"]:
            if item["scheme_code"].upper() == self.agreement.upper():
                # if item["code"].upper() == self.agreement.upper():
                self.rule_offset = item["rule_offset"]
                self.country_prefix = item["scheme_code"]
                self.country_code = item["country_code"]
                break

    def create(self):
        self.document = docx.Document(self.source)
        table = self.document.tables[0]
        table_cells = []
        table_cells2 = []

        keys = None
        for i, row in enumerate(table.rows):
            table_cell = TableCell(self.rule_offset, self.fetch_descriptions, self.country_code, self.country_prefix, i + 1)

            table_cell.cell_classification = row.cells[0].text
            table_cell.cell_classification = table_cell.cell_classification.replace("\n", "")
            table_cell.cell_description = row.cells[1].text
            table_cell.cell_specific = row.cells[2].text
            table_cell.cell_psr = row.cells[3].text

            if table_cell.cell_classification != "Classification":
                table_cell.parse()
                if table_cell.valid:
                    del table_cell.valid
                    table_cells.append(table_cell)
                    self.rule_offset += 1

            if i > self.max_row_count:
                 break

        # Sort the rules by low code
        table_cells.sort(key=lambda x: x.key_first, reverse=False)

        # Fill in the missing key_last values
        for i in range(0, len(table_cells) - 1):
            tc = table_cells[i]
            tc2 = table_cells[i]
            if "09.01" in tc.cell_classification:
                a = 1
            if tc.key_last is None:
                if tc.key_first[-6:] == "000000":
                    tc.key_last = tc.key_first[0:4] + "999999"
                elif tc.key_first[-4:] == "0000":
                    tc.key_last = tc.key_first[0:6] + "9999"
                else:
                    try:
                        tc.key_last = str(int(tc2.key_first) - 1).ljust(10, "0")
                    except:
                        tc.key_last = None

        self.delete_rows_from_roo_table()

        # Create the JSON object
        for tc in table_cells:
            tc.pick_out_key_terms()
            tc.remove_unnecessary_fields()
            tc.write_table_cell_to_db()
            table_cells2.append(tc.__dict__)

        # This takes the data that has been assigned in the 'rules'
        # table and assigns it to the rules_to_commodities table
        self.write_to_commodities_table(table_cells)
        self.export_to_json()

        # All the code below is not used anymore
        # This created data in a format that was initially used for the UK codes, but is no longer valid

        if 1 > 2:
            json_string = json.dumps(table_cells2, indent=2)
            f = open(self.dest, "w+")
            f.write(json_string)
            f.close()

            self.dest_filename = os.path.join(
                self.roo_folder, self.agreement.lower() + ".json")
            shutil.copy(self.dest, self.dest_filename)

    def export_to_json(self):
        d = Database()

        sql = """
        select sub_heading, heading, description, rule,
        alternate_rule, quota_amount, quota_unit, key_first, key_last, r.id_rule
        from roo.rules_to_commodities rtc, roo.rules r
        where r.id_rule = rtc.id_rule 
        and r.country_prefix = %s
        and r.scope = %s
        and rtc.scope = %s
        order by sub_heading, id_rule;
        """

        params = [
            self.agreement.lower(),
            "uk",
            "uk"
        ]
        rows = d.run_query(sql, params)
        object = {}
        object["rules"] = []
        previous_rule = Rule()
        for row in rows:
            rule = Rule(row)
            object["rules"].append(rule.asdict())
            previous_rule = rule

        with open(self.dest_filename, 'w') as f:
            json.dump(object, f, indent=4)
        f.close()

    def write_to_commodities_table(self, table_cells):
        self.delete_from_commodities_table()
        # Take the key_first - 1st 6 digits
        # Take the key_last - 1st 6 digits
        # Assign the rule to each code from the exemplar codes, where the key_first and key_last fit into the exemplar_code (1st 6 digits)
        for tc in table_cells:
            if tc.cell_classification == "200819":
                a = 1
            print("Writing row ", tc.cell_classification)
            for ex in self.exemplar_codes:
                if tc.key_last is None:
                    tc.key_last = tc.key_first[0:6] + "9999"
                    a = 1
                    
                if tc.key_first[0:6] <= ex.hs_code and tc.key_last[0:6] >= ex.hs_code:
                    a = 1
                    d = Database()
                    sql = """
                    insert into roo.rules_to_commodities
                    (
                        id_rule,
                        sub_heading,
                        country_prefix,
                        scope
                    ) values (%s, %s, %s, %s)
                    """
                    params = [
                        tc.id,
                        ex.hs_code,
                        self.agreement.lower(),
                        "uk"
                    ]
                    d.run_query(sql, params)
                try:
                    if ex.hs_code > tc.key_last:
                        break
                except:
                    a = 1
            a = 1

    def delete_rows_from_roo_table(self):
        print("\nDeleting rows from RoO table\n")
        d = Database()
        sql = """
        delete from roo.rules
        where scope = 'uk' and lower(country_prefix) = %s
        """
        params = [
            self.country_prefix.lower()
        ]
        d.run_query(sql, params)

    def delete_from_commodities_table(self):
        print("\nDeleting rows from RoO to commodities table\n")
        d = Database()
        sql = """
        delete from roo.rules_to_commodities
        where scope = 'uk' and lower(country_prefix) = %s
        """
        params = [
            self.country_prefix.lower()
        ]
        d.run_query(sql, params)
